# -*- coding: utf-8 -*-
import re
from pathlib import Path
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH

INVALID = r'<>:"/\\|?*'

def safe_name(name: str, maxlen: int = 120) -> str:
    name = name.strip()
    for ch in INVALID:
        name = name.replace(ch, "_")
    name = re.sub(r'\s+', ' ', name).strip()
    if len(name) > maxlen:
        name = name[:maxlen].rstrip()
    return name or "chapter"

def save_chapter_docx(folder: Path, chapter_title: str, text: str, index: int = None):
    folder.mkdir(parents=True, exist_ok=True)
    base = safe_name(chapter_title)
    filename = f"{index:03d} {base}.docx" if index is not None else f"{base}.docx"
    path = folder / filename
    doc = Document()
    h = doc.add_heading(chapter_title, level=1)
    h.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for line in text.split("\n"):
        if line.strip():
            doc.add_paragraph(line.strip())
        else:
            doc.add_paragraph()
    doc.save(path)
    return path
